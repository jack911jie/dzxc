import os
import sys
import pandas as pd
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式


class Docx:
    def __init__(self):
        pass

    def read_excel(self,xls='E:\\WXWork\\1688856932305542\\WeDrive\\大智小超科学实验室\\2-乐高课程\\课程信息表.xlsx'):
        df=pd.read_excel(xls)
        mask = [ True if i[0] == 'L' else False for i in df['课程编号'] ]
        # df=df[df.课程编号.str.startwith('L')]
        df=df[mask]
        df['id_crs']=df['课程编号']+df['课程名称']
        
        crs=df['id_crs'].tolist()
        crs_txt=df['课程描述'].tolist()
        pyq=df['朋友圈文案'].tolist()

        info=[]
        for n,crss in enumerate(crs):
            info.append([crss,crs_txt[n],pyq[n]])

        # print(len(info))
        return info

    def create_word(self):
        info=self.read_excel()
        
        doc=Document()
        for inf in info:
            self.word_texts(doc,inf[0],inf[1],inf[2])

        doc.save("e:\\temp\\temp_dzxc\\测试文件.docx")


    def word_texts(self,doc,title,crs_txt,pyq):
        #创建一个空白的word文档
        

        #设置1级标题
        para_heading_1=doc.add_heading('',level=1)#返回1级标题段落对象，标题也相当于一个段落
        # # para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
        # # para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        # # para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        # # para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
        # # para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        # # para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸
        run=para_heading_1.add_run(title)
        # run.font.name=u'宋体'    #设置为宋体
        # # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # # run.font.size=Pt(12)#设置1级标题文字的大小为“小四” 为12磅
        run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

        #设置2级标题
        # para_heading_2=doc.add_heading('',level=2)#返回1级标题段落对象，标题也相当于一个段落
        # # para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
        # # para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        # # para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        # # para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
        # # para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        # # para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸
        # run=para_heading_2.add_run("课程描述")
        # run.font.name=u'宋体'    #设置为宋体
        # # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # # run.font.size=Pt(12)#设置1级标题文字的大小为“小四” 为12磅
        # run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

        #######################课程描述标题##########################
        #增加一段文字
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        p.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        p.paragraph_format.line_spacing=1.5  #设置行间距为 1.5倍
        #p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
        p.paragraph_format.first_line_indent=Inches(0.3346457)#相当于小四两个字符的缩进

        p.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        p.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

        r=p.add_run('课程描述')
        # r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # r.font.size=Pt(12)  #设置字体大小为12磅 相当于 小四
        # r.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色 
        r.bold=True

        #######################课程描述##########################
        #增加一段文字
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        p.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        p.paragraph_format.line_spacing=1.5  #设置行间距为 1.5倍
        #p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
        p.paragraph_format.first_line_indent=Inches(0.3346457)#相当于小四两个字符的缩进

        p.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        p.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

        r=p.add_run(crs_txt)
        # r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # r.font.size=Pt(12)  #设置字体大小为12磅 相当于 小四
        # r.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色 




        # #####朋友圈文案####
        # #设置2级标题
        # para_heading_2=doc.add_heading('',level=2)#返回1级标题段落对象，标题也相当于一个段落
        # # para_heading.alignment=WD_ALIGN_PARAGRAPH.LEFT#设置为左对齐
        # # para_heading.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        # # para_heading.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        # # para_heading.paragraph_format.line_spacing=1.5 #设置行间距为 1.5
        # # para_heading.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        # # para_heading.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸
        # run=para_heading_2.add_run("朋友圈文案")
        # run.font.name=u'宋体'    #设置为宋体
        # # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # # run.font.size=Pt(12)#设置1级标题文字的大小为“小四” 为12磅
        # run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

         #######################朋友圈文案标题##########################
        #增加一段文字
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        p.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        p.paragraph_format.line_spacing=1.5  #设置行间距为 1.5倍
        #p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
        p.paragraph_format.first_line_indent=Inches(0.3346457)#相当于小四两个字符的缩进

        p.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        p.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

        r=p.add_run('朋友圈文案')
        # r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # r.font.size=Pt(12)  #设置字体大小为12磅 相当于 小四
        # r.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色 
        r.bold=True

        #增加一段文字，朋友圈文案正文
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)#设置段前 0 磅
        p.paragraph_format.space_after=Pt(0) #设置段后 0 磅
        p.paragraph_format.line_spacing=1.5  #设置行间距为 1.5倍
        #p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
        p.paragraph_format.first_line_indent=Inches(0.3346457)#相当于小四两个字符的缩进

        p.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
        p.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸

        r=p.add_run(pyq)
        # r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
        # r.font.size=Pt(12)  #设置字体大小为12磅 相当于 小四
        # r.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色 
        

        


if __name__=='__main__':
    ppp=Docx()
    ppp.create_word()
    # ppp.read_excel()